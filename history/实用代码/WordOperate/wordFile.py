'''
升级pip：python -m pip install -U pip setuptools

安装库：
python-docx    （大部分操作都是使用此库）
win32com    （主要用作doc转docx格式转换用）
mailmerge    （用作按照模板生成大量同类型文档）
matplotlib    （Python的绘图库，不深入了解）
'''

#导入包
from docx import Document
#新建，复制word文档
def GenerateNewWord(fileName):
    #可添加一个判断末尾是否是 .docx 的函数
    #fileName = fileName+'.docx'
    print('即将新建文件：',fileName)
    #新建文件
    document = Document()
    document.save(fileName)
    #新建指定文件名  **  document.save('exist.docx')

#复制已存在 .docx 的文件
def CopyWord(fileName):
    document = Document(fileName)
    fileName = fileName.replace('.docx','(复制).docx')
    print('复制文件名为：',fileName)
    document.save(fileName)
    '''
    document = Document('exist.docx')
    document.save('new.docx')
    '''


#win32com将doc转为docx
from win32com import client as wc
def TransDocToDocx(oldDocName,newDocxName):
    #新建文件，注：利用线程的方式将文件转换变为线程程序存放于计算机中，若有文件存在，则将其转换为。docx格式
    document = Document()
    document.save('旧doc格式文档.doc')

    print("正在将文件"+oldDocName+"转换为"+newDocxName)
    #打开word
    word = wc.Dispatch('Word.Application')

    #打开旧word文件
    doc = word.Documents.Open(oldDocName)

    #保存为新word文件，其中参数12表示的是docx文件
    doc.SaveAs(newDocxName,12)

    #关闭word文档
    doc.Close()
    word.Quit()

    print("文件转换完毕！")


def funOpenExistFile():
    word = Dispatch('Word.Application')
    #或者可以使用下面的方法，使用启动独立的进程
    # word = DisPatchEx('Word.Application')

    #如果不声明以下属性，运行的时候会显示打开的word
    word.Visible = 1 # 0:后台运行 1:前台运行（可见）
    word.DisplayAlerts = 0 #不显示，不警告

    # 创建新的word文档
    doc = word.Documents.Add()

    #在文档开头添加内容
    myRange1 = doc.Range(0,0)
    myRange1.InsertBefore('Hello wordn'+'\n')

    #在文档末尾添加内容
    myRange2 = doc.Range()
    myRange2.InsertAfter('Bye wordn')

    #在文档i指定位置添加内容
    i=12
    myRange3 = doc.Range(0,i)
    myRange3.InsertAfter('what\'s up, bro?n'+'\n')

    #doc.Sava() #保存
    doc.SaveAs(os.getcwd() + 'funOpenNewFile.docx')
    doc.Close()  #关闭word文档
    word.Quit()  #关闭office


#将word转换为pdf
from win32com.client import Dispatch
import os
import time
#生成pdf文件
def funGeneratePDF():
    word = Dispatch('Word.Application')
    #word.Visible = 0 #后台运行，不显示
    word.Visible = 1  # 0:后台运行 1:前台运行（可见）
    word.DisplayAlerts = 0  # 不显示，不警告
    currentPath = os.getcwd()

    print(os.path.join(currentPath,'JavaWeb期末文件_lv.docx'))
    doc = word.Documents.Open(os.path.join(currentPath,'JavaWeb期末文件_lv.docx')) #打开一个已有的word文档
    time.sleep(2)

    doc.SaveAs(os.path.join(currentPath,'JavaWeb期末文件_lv.pdf'),17)  # txt = 4 ,html = 10,docx = 16, pdf = 17
    doc.Close()
    word.Quit()




if __name__ == '__main__':
    # #用menu或以类的形式创建菜单目录
    # print('记得带上文件末尾，否则将无法正确完成文件创建操作')
    # #新建 .docx 文件
    # GenerateNewWord('函数新建文件.docx')
    # #复制 .docx文件
    # CopyWord('函数新建文件.docx')
    #
    # #新建一个 .doc 文件以便进行文件转换
    # GenerateNewWord('旧doc格式文档.doc')
    # #获取当前目录完整路径
    # currentPath = os.getcwd()
    # print("当前路径为：",currentPath)
    #
    # #获取就doc格式word文件绝对路径名
    # docName = os.path.join(currentPath,'旧doc格式文档.doc')
    # print("docFilePath = ",docName)
    #
    # #设置新docx格式文档文件名
    # docxName = os.path.join(currentPath,'新生成docx格式文档.docx')
    #
    # TransDocToDocx(docName,docxName)
    # #新建一个文档，并对其进行写入操作，个人感觉无用，除非是进行多文件写入操作
    # funOpenExistFile()

    funGeneratePDF()


